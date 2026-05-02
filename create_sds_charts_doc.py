"""Generate SDS_Charts_and_Residuals.docx — reference tables for Tom."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Formatting helpers (same palette as Getting Started Guide) ──────────


def set_cell_font(cell, size=Pt(9), bold=False, font_name="Calibri"):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = size
            run.bold = bold
            run.font.name = font_name


def shade_row(row, fill="D9E2F3"):
    """Apply background shading to every cell in a row."""
    for cell in row.cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill,
        })
        tc_pr.append(shd)


def add_table_from_data(doc, headers, rows, col_widths=None):
    """Create a table with header row + data rows.  Returns the table object."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        hdr.cells[i].text = text
    shade_row(hdr, fill="4472C4")
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
            set_cell_font(row.cells[c_idx], size=Pt(9))
        if r_idx % 2 == 1:
            shade_row(row, fill="E9EDF4")

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Inches(w)

    return table


def add_footer_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


# ── Document ────────────────────────────────────────────────────────────


def build_document():
    doc = Document()

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title
    title = doc.add_heading("SDS Charts & Residuals Reference", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Quick-reference tables for the processbehavior library. "
        "Use the Notes column for annotations."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Table 1 — Primary Charts by SDS ─────────────────────────────────
    doc.add_heading("Table 1 — Primary Charts by SDS", level=1)

    t1_headers = [
        "SDS", "Name", "Grid", "Replication",
        "Xbar", "S", "XmR", "R", "Histogram", "Recommended", "Notes",
    ]
    Y = "Y"
    t1_rows = [
        [1, "Full Factorial, Complete Rep.", "Complete", "All n >= 2",
         Y, Y, Y, Y, Y, "Xbar", ""],
        [2, "Full Factorial, No Rep.", "Complete", "All n = 1",
         Y, Y, Y, Y, Y, "Xbar", ""],
        [3, "Partial Replication", "Complete", "Mixed",
         Y, Y, Y, Y, Y, "Xbar", ""],
        [4, "Incomplete + Singletons", "Incomplete", "Mixed",
         Y, Y, Y, Y, Y, "Xbar", ""],
        [5, "Incomplete, No Singletons", "Incomplete", "All n >= 2",
         Y, Y, Y, Y, Y, "Xbar", ""],
        [6, "Incomplete, No Replication", "Incomplete", "All n <= 1",
         "\u2014", "\u2014", Y, Y, Y, "XmR", ""],
    ]
    add_table_from_data(doc, t1_headers, t1_rows)
    add_footer_note(
        doc,
        "SDS 6 cannot use Xbar or S \u2014 no within-cell variance.",
    )

    # ── Table 2 — VAS Residuals by SDS ──────────────────────────────────
    doc.add_heading("Table 2 — VAS Residuals by SDS", level=1)

    t2_headers = [
        "SDS", "R1", "R2", "R3", "R4", "R5",
        "R2 Method", "VAS Computed?", "Notes",
    ]
    t2_rows = [
        [1, Y, Y, Y, Y, Y, "exact",          "Yes (Xbar/S)", ""],
        [2, Y, Y, Y, Y, Y, "moving_average",  "Yes (Xbar/S)", ""],
        [3, Y, Y, Y, Y, Y, "hybrid",          "Yes (Xbar/S)", ""],
        [4, Y, Y, Y, Y, Y, "hybrid",          "Yes (Xbar/S)", ""],
        [5, Y, Y, Y, Y, Y, "exact",           "Yes (Xbar/S)", ""],
        [6, Y, Y, Y, Y, Y, "moving_average",  "No*",          ""],
    ]
    add_table_from_data(doc, t2_headers, t2_rows)
    add_footer_note(
        doc,
        "*SDS 6 sets vas_residuals_supported=True in the plan, but "
        "should_calculate_vas_residuals() returns False unconditionally "
        "for SDS 6 (line 820). Since SDS 6 only allows XmR/R/Histogram, "
        "the XmR/R exclusion also blocks it.",
    )

    # ── Table 3 — R2 Method Rules ───────────────────────────────────────
    doc.add_heading("Table 3 — R2 Method Rules", level=1)

    t3_headers = [
        "Condition", "R2 Method", "Formula",
        "Wheeler Eq", "SDS Examples", "Notes",
    ]
    t3_rows = [
        ["All cells n >= 2", "exact",
         "R2 = Y \u2212 \u0232_kt", "59", "1, 5", ""],
        ["All cells n = 1", "ma2",
         "R2 = (Y_j \u2212 Y_{j\u22121}) / 2", "66", "2, 6", ""],
        ["Mixed", "hybrid",
         "exact where n>=2, ma2 where n=1", "\u2014", "3, 4", ""],
    ]
    add_table_from_data(doc, t3_headers, t3_rows)

    # ── Table 4 — Residual Chart Types by SDS ───────────────────────────
    doc.add_heading("Table 4 — Residual Chart Types by SDS", level=1)
    doc.add_paragraph(
        "These use the old Residual_ChartType naming "
        "(R2_S, R3_Xbar, etc.) \u2014 see GitHub issue for cleanup proposal."
    )

    t4_headers = [
        "Residual", "Selection Rule",
        "SDS 1", "SDS 2", "SDS 3", "SDS 4", "SDS 5", "Notes",
    ]
    t4_rows = [
        ["R2", "S if min_cell>=2, else XmR",
         "R2_S", "R2_XmR", "runtime", "runtime", "R2_S", ""],
        ["R3", "Xbar+S if min_cell>=2, else XmR",
         "R3_Xbar, R3_S", "R3_XmR", "runtime", "runtime",
         "R3_Xbar, R3_S", ""],
        ["R4", "Xbar+S if has_factors, else XmR",
         "R4_Xbar, R4_S", "R4_Xbar, R4_S", "R4_Xbar, R4_S",
         "R4_Xbar, R4_S", "R4_Xbar, R4_S", ""],
        ["R5", "Xbar+S if has_time, else XmR",
         "R5_Xbar, R5_S", "R5_Xbar, R5_S", "R5_Xbar, R5_S",
         "R5_Xbar, R5_S", "R5_Xbar, R5_S", ""],
    ]
    add_table_from_data(doc, t4_headers, t4_rows)
    add_footer_note(
        doc,
        "SDS 6 excluded \u2014 residual_charts returns [].",
    )

    # ── Table 5 — Effects & Capabilities by SDS ─────────────────────────
    doc.add_heading("Table 5 — Effects & Capabilities by SDS", level=1)

    t5_headers = [
        "SDS", "Main Effects", "Interactions",
        "Stratification", "Notes",
    ]
    t5_rows = [
        [1, "Yes", "Yes", "Yes", ""],
        [2, "Yes", "Yes", "Yes", ""],
        [3, "Yes", "Yes", "Yes", ""],
        [4, "Yes", "No",  "Yes", ""],
        [5, "Yes", "No",  "Yes", ""],
        [6, "No",  "No",  "Yes", ""],
    ]
    add_table_from_data(doc, t5_headers, t5_rows)

    # ── Table 6 — Residual Definitions ──────────────────────────────────
    doc.add_heading("Table 6 — Residual Definitions", level=1)

    t6_headers = [
        "Residual", "Measures", "Formula", "Wheeler Eq", "Notes",
    ]
    t6_rows = [
        ["R1", "Total deviation from grand mean",
         "Y \u2212 \u0232", "56", ""],
        ["R2", "Within-cell (unexplained)",
         "varies by method", "59/66", ""],
        ["R3", "Factor \u00d7 Time interaction",
         "Y \u2212 \u0232_k \u2212 \u0232_t + \u0232", "66", ""],
        ["R4", "Time effects + unexplained",
         "(\u0232_t \u2212 \u0232) + R2", "72", ""],
        ["R5", "Factor effects + unexplained",
         "(\u0232_k \u2212 \u0232) + R2", "75", ""],
    ]
    add_table_from_data(doc, t6_headers, t6_rows)

    return doc


if __name__ == "__main__":
    doc = build_document()
    out = "processbehavior/SDS_Charts_and_Residuals.docx"
    doc.save(out)
    print(f"Wrote {out}")
